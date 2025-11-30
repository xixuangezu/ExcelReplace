# -*- coding: utf-8 -*- 
import os, re, configparser, tkinter as tk
from tkinter import filedialog, scrolledtext, messagebox
from openpyxl import load_workbook
from docx import Document
import sys

class ReplaceTool(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("批量替换工具")
        self.geometry("900x600")
        if getattr(sys, 'frozen', False):
            base_dir = os.path.dirname(sys.executable)
        else:
            base_dir = os.path.dirname(os.path.abspath(__file__))
        self.config_file = os.path.join(base_dir, ".replace_tool.ini")
        # 选择按钮 + 路径显示
        tk.Button(self, text="选择文件/目录", command=self.select_target).grid(row=0, column=0, padx=5, pady=5)
        self.path_var = tk.StringVar()
        tk.Entry(self, textvariable=self.path_var, width=100).grid(row=0, column=1, padx=5, pady=5)
        # 读取上次记录
        cfg = configparser.ConfigParser()
        if os.path.exists(self.config_file):
            cfg.read(self.config_file)
            if cfg.has_option("settings", "last_target"):
                self.path_var.set(cfg.get("settings", "last_target"))
        # 查找/替换/正则
        tk.Label(self, text="查找：").grid(row=1, column=0, sticky="e")
        self.find_var = tk.StringVar()
        tk.Entry(self, textvariable=self.find_var, width=80).grid(row=1, column=1, sticky="w", padx=(5,0))
        tk.Label(self, text="替换：").grid(row=2, column=0, sticky="e")
        self.replace_var = tk.StringVar()
        tk.Entry(self, textvariable=self.replace_var, width=80).grid(row=2, column=1, sticky="w", padx=(5,0))
        self.use_regex = tk.BooleanVar()
        tk.Checkbutton(self, text="正则匹配", variable=self.use_regex).grid(row=1, column=2, rowspan=2, padx=5)
        # 执行按钮
        tk.Button(self, text="预览", command=self.run_preview).grid(row=3, column=0, pady=5)
        tk.Button(self, text="开始替换", command=self.run_replace).grid(row=3, column=1, pady=5, sticky="w")
        # 日志窗口
        self.log = scrolledtext.ScrolledText(self, wrap=tk.WORD)
        self.log.tag_config("before", foreground="red")
        self.log.tag_config("after", foreground="green")
        self.log.grid(row=4, column=0, columnspan=3, sticky="nsew", padx=5, pady=5)
        self.grid_rowconfigure(4, weight=1)
        self.grid_columnconfigure(1, weight=1)

    def save_config(self):
        cfg = configparser.ConfigParser()
        cfg["settings"] = {"last_target": self.path_var.get()}
        with open(self.config_file, "w") as f:
            cfg.write(f)

    def select_target(self):
        menu = tk.Menu(self, tearoff=0)
        menu.add_command(label="目录", command=lambda:self._select("dir"))
        menu.add_command(label="文件", command=lambda:self._select("file"))
        menu.post(self.winfo_pointerx(), self.winfo_pointery())

    def _select(self, mode):
        init = self.path_var.get() or os.getcwd()
        if mode == "dir":
            d = filedialog.askdirectory(initialdir=init)
            if d:
                self.path_var.set(d); self.save_config()
        else:
            fs = filedialog.askopenfilenames(
                initialdir=init,
                filetypes=[("Excel/.xlsx","*.xlsx"),("Word/.docx","*.docx")]
            )
            if fs:
                self.path_var.set(";".join(fs)); self.save_config()

    def run_preview(self):
        self._run(replace=False)

    def run_replace(self):
        self._run(replace=True)

    def _run(self, replace):
        target = self.path_var.get().strip()
        if not target:
            messagebox.showwarning("警告", "请先选择目录或文件！"); return
        self.log.delete("1.0", tk.END)
        paths = []
        if os.path.isdir(target):
            base = target
            for r, _, files in os.walk(target):
                for f in files:
                    if f.lower().endswith((".xlsx",".docx")):
                        paths.append(os.path.join(r, f))
        else:
            base = os.getcwd()
            for p in target.split(";"):
                if os.path.isfile(p) and p.lower().endswith((".xlsx",".docx")):
                    paths.append(p)
        if not paths:
            messagebox.showinfo("提示", "没找到 .xlsx/.docx 文件。"); return
        pat, rep, use_re = self.find_var.get(), self.replace_var.get(), self.use_regex.get()
        for p in paths:
            rel = os.path.relpath(p, base)
            idx = self.log.index(tk.END)
            if p.lower().endswith(".xlsx"):
                changed = self._process_xlsx(p, pat, rep, use_re, replace)
            else:
                changed = self._process_docx(p, pat, rep, use_re, replace)
            if changed:
                self.log.insert(idx, f"## {rel}\n")
        if replace:
            messagebox.showinfo("完成", "替换完成" )

    def _process_xlsx(self, fn, pat, rep, use_re, do_replace):
        wb = load_workbook(fn); changed=False
        for name in wb.sheetnames:
            ws=wb[name]
            for row in ws.iter_rows():
                for cell in row:
                    v=cell.value
                    if isinstance(v,str):
                        new = re.sub(pat,rep,v) if use_re else v.replace(pat,rep)
                        if new!=v:
                            changed=True
                            self.log.insert(tk.END, f"  {cell.coordinate} ")
                            self.log.insert(tk.END, v, "before")
                            self.log.insert(tk.END, " -> ")
                            self.log.insert(tk.END, new+"\n","after")
                            if do_replace: cell.value=new
        if do_replace and changed: wb.save(fn)
        return changed

    def _process_docx(self, fn, pat, rep, use_re, do_replace):
        doc=Document(fn); changed=False
        for para in doc.paragraphs:
            full=para.text
            new = re.sub(pat,rep,full) if use_re else full.replace(pat,rep)
            if new!=full:
                changed=True
                self.log.insert(tk.END, "  段落 ")
                self.log.insert(tk.END, full,"before")
                self.log.insert(tk.END, " -> ")
                self.log.insert(tk.END, new+"\n","after")
                if do_replace:
                    for run in para.runs: run.text=""
                    para.add_run(new)
        if do_replace and changed: doc.save(fn)
        return changed

if __name__ == "__main__":
    ReplaceTool().mainloop()
